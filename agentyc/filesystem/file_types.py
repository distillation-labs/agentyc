import asyncio
import csv
import io
from abc import ABC, abstractmethod
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path

from pydantic import BaseModel

from agentyc.filesystem.errors import FileSystemError


class BaseFile(BaseModel, ABC):
	"""Base class for all file types."""

	name: str
	content: str = ''

	@property
	@abstractmethod
	def extension(self) -> str:
		"""File extension (e.g. 'txt', 'md')."""

	def write_file_content(self, content: str) -> None:
		self.update_content(content)

	def append_file_content(self, content: str) -> None:
		self.update_content(self.content + content)

	def update_content(self, content: str) -> None:
		self.content = content

	def sync_to_disk_sync(self, path: Path) -> None:
		(path / self.full_name).write_text(self.content)

	async def sync_to_disk(self, path: Path) -> None:
		file_path = path / self.full_name
		with ThreadPoolExecutor() as executor:
			await asyncio.get_event_loop().run_in_executor(executor, lambda: file_path.write_text(self.content))

	async def write(self, content: str, path: Path) -> None:
		self.write_file_content(content)
		await self.sync_to_disk(path)

	async def append(self, content: str, path: Path) -> None:
		self.append_file_content(content)
		await self.sync_to_disk(path)

	def read(self) -> str:
		return self.content

	@property
	def full_name(self) -> str:
		return f'{self.name}.{self.extension}'

	@property
	def get_size(self) -> int:
		return len(self.content)

	@property
	def get_line_count(self) -> int:
		return len(self.content.splitlines())


class MarkdownFile(BaseFile):
	@property
	def extension(self) -> str:
		return 'md'


class TxtFile(BaseFile):
	@property
	def extension(self) -> str:
		return 'txt'


class JsonFile(BaseFile):
	@property
	def extension(self) -> str:
		return 'json'


class CsvFile(BaseFile):
	"""CSV file implementation with automatic RFC 4180 normalization."""

	@property
	def extension(self) -> str:
		return 'csv'

	@staticmethod
	def _normalize_csv(raw: str) -> str:
		stripped = raw.strip('\n\r')
		if not stripped:
			return raw

		if '\n' not in stripped and '\\n' in stripped:
			stripped = stripped.replace('\\"', '"')
			stripped = stripped.replace('\\n', '\n')

		reader = csv.reader(io.StringIO(stripped))
		rows: list[list[str]] = []
		for row in reader:
			if row:
				rows.append(row)

		if not rows:
			return raw

		out = io.StringIO()
		writer = csv.writer(out, lineterminator='\n')
		writer.writerows(rows)
		return out.getvalue().rstrip('\n')

	def write_file_content(self, content: str) -> None:
		self.update_content(self._normalize_csv(content))

	def append_file_content(self, content: str) -> None:
		normalized_new = self._normalize_csv(content)
		if not normalized_new.strip('\n\r'):
			return
		existing = self.content
		if existing and not existing.endswith('\n'):
			existing += '\n'
		self.update_content(self._normalize_csv(existing + normalized_new))


class JsonlFile(BaseFile):
	@property
	def extension(self) -> str:
		return 'jsonl'


class PdfFile(BaseFile):
	@property
	def extension(self) -> str:
		return 'pdf'

	def sync_to_disk_sync(self, path: Path) -> None:
		from reportlab.lib.pagesizes import letter
		from reportlab.lib.styles import getSampleStyleSheet
		from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

		file_path = path / self.full_name
		try:
			doc = SimpleDocTemplate(str(file_path), pagesize=letter)
			styles = getSampleStyleSheet()
			story = []
			for line in self.content.split('\n'):
				if line.strip():
					if line.startswith('# '):
						para = Paragraph(line[2:], styles['Title'])
					elif line.startswith('## '):
						para = Paragraph(line[3:], styles['Heading1'])
					elif line.startswith('### '):
						para = Paragraph(line[4:], styles['Heading2'])
					else:
						para = Paragraph(line, styles['Normal'])
					story.append(para)
				else:
					story.append(Spacer(1, 6))
			doc.build(story)
		except Exception as e:
			raise FileSystemError(f"Error: Could not write to file '{self.full_name}'. {str(e)}")

	async def sync_to_disk(self, path: Path) -> None:
		with ThreadPoolExecutor() as executor:
			await asyncio.get_event_loop().run_in_executor(executor, lambda: self.sync_to_disk_sync(path))


class DocxFile(BaseFile):
	@property
	def extension(self) -> str:
		return 'docx'

	def sync_to_disk_sync(self, path: Path) -> None:
		file_path = path / self.full_name
		try:
			from docx import Document

			doc = Document()
			for line in self.content.split('\n'):
				if line.strip():
					if line.startswith('# '):
						doc.add_heading(line[2:], level=1)
					elif line.startswith('## '):
						doc.add_heading(line[3:], level=2)
					elif line.startswith('### '):
						doc.add_heading(line[4:], level=3)
					else:
						doc.add_paragraph(line)
				else:
					doc.add_paragraph()
			doc.save(str(file_path))
		except Exception as e:
			raise FileSystemError(f"Error: Could not write to file '{self.full_name}'. {str(e)}")

	async def sync_to_disk(self, path: Path) -> None:
		with ThreadPoolExecutor() as executor:
			await asyncio.get_event_loop().run_in_executor(executor, lambda: self.sync_to_disk_sync(path))


class HtmlFile(BaseFile):
	@property
	def extension(self) -> str:
		return 'html'


class XmlFile(BaseFile):
	@property
	def extension(self) -> str:
		return 'xml'


FILE_TYPE_CLASSES: dict[str, type[BaseFile]] = {
	'md': MarkdownFile,
	'txt': TxtFile,
	'json': JsonFile,
	'jsonl': JsonlFile,
	'csv': CsvFile,
	'pdf': PdfFile,
	'docx': DocxFile,
	'html': HtmlFile,
	'xml': XmlFile,
}

FILE_TYPE_NAME_MAP: dict[str, type[BaseFile]] = {file_type.__name__: file_type for file_type in FILE_TYPE_CLASSES.values()}
